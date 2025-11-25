"""
Модуль для извлечения текста из документов тендерной документации.
Поддерживает форматы: PDF, DOCX, XLSX, RTF, TXT, CSV, ZIP и любые текстовые файлы.
"""

import os
from typing import Optional
from pathlib import Path
import PyPDF2
from docx import Document
import subprocess
import zipfile
import tempfile
import shutil
import csv
import chardet
from openpyxl import load_workbook


class TextExtractor:
    """Класс для извлечения текста из различных форматов документов."""

    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """
        Определяет реальный тип файла по содержимому (magic bytes), а не по расширению.
        КРИТИЧНО для zakupki.gov.ru которые часто называют DOCX файлы как .pdf

        Args:
            file_path: Путь к файлу

        Returns:
            Тип файла: 'pdf', 'docx', 'doc', 'xlsx', 'rtf', 'txt', 'csv', 'text' или 'unknown'
        """
        try:
            # МЕТОД 1: Читаем magic bytes напрямую (самый надежный)
            with open(file_path, 'rb') as f:
                magic = f.read(8)

            # Логируем magic bytes для отладки
            magic_hex = ' '.join(f'{b:02X}' for b in magic[:4])
            print(f"   🔬 Magic bytes: {magic_hex}")

            # ZIP-based форматы (DOCX, XLSX) начинаются с PK (50 4B)
            if magic[:2] == b'PK':
                # Это ZIP архив - может быть DOCX или XLSX
                # Читаем дальше чтобы определить точный тип
                try:
                    with zipfile.ZipFile(file_path, 'r') as zip_ref:
                        namelist = zip_ref.namelist()
                        # DOCX содержит word/document.xml
                        if any('word/' in name for name in namelist):
                            print(f"   ✅ Определен как DOCX (найдено word/ в архиве)")
                            return 'docx'
                        # XLSX содержит xl/
                        elif any('xl/' in name for name in namelist):
                            print(f"   ✅ Определен как XLSX (найдено xl/ в архиве)")
                            return 'xlsx'
                        else:
                            print(f"   ⚠️  ZIP архив неизвестного типа")
                            return 'zip'
                except Exception as e:
                    print(f"   ⚠️  Ошибка чтения ZIP: {e}")
                    return 'zip'

            # PDF начинается с %PDF (25 50 44 46)
            elif magic[:4] == b'%PDF':
                print(f"   ✅ Определен как PDF (magic bytes: %PDF)")
                return 'pdf'

            # RTF начинается с {\rtf
            elif magic[:5] == b'{\\rtf':
                print(f"   ✅ Определен как RTF")
                return 'rtf'

            # Старые DOC файлы (OLE Compound Document) начинаются с D0 CF 11 E0
            # НО! Это также может быть DOCX или XLSX, сохранённый в старом формате
            # Проверим через zipfile - если это ZIP, значит это современный Office
            elif magic[:4] == b'\xD0\xCF\x11\xE0':
                print(f"   🔍 OLE Compound Document обнаружен")
                # Пытаемся открыть как ZIP (для DOCX/XLSX)
                try:
                    with zipfile.ZipFile(file_path, 'r') as zip_ref:
                        namelist = zip_ref.namelist()
                        if any('word/' in name for name in namelist):
                            print(f"   ✅ Определен как DOCX (в OLE контейнере)")
                            return 'docx'
                        elif any('xl/' in name for name in namelist):
                            print(f"   ✅ Определен как XLSX (в OLE контейнере)")
                            return 'xlsx'
                except:
                    pass  # Не ZIP, значит реально старый DOC

                print(f"   ✅ Определен как старый DOC")
                return 'doc'

            # МЕТОД 2: Пробуем системную команду 'file' (если доступна)
            try:
                result = subprocess.run(
                    ['file', '--brief', '--mime-type', file_path],
                    capture_output=True,
                    text=True,
                    timeout=5
                )

                if result.returncode == 0:
                    mime_type = result.stdout.strip().lower()

                    # Определяем тип по MIME
                    if 'pdf' in mime_type:
                        return 'pdf'
                    elif 'spreadsheetml' in mime_type or 'ms-excel' in mime_type:
                        return 'xlsx'
                    elif 'wordprocessingml' in mime_type or 'vnd.openxmlformats' in mime_type:
                        return 'docx'
                    elif 'msword' in mime_type or 'ms-word' in mime_type:
                        return 'doc'
                    elif 'rtf' in mime_type or 'richtext' in mime_type:
                        return 'rtf'
                    elif 'csv' in mime_type or 'comma-separated' in mime_type:
                        return 'csv'
                    elif 'text/plain' in mime_type or 'text/' in mime_type:
                        return 'txt'
                    elif 'composite' in mime_type or 'ole' in mime_type:
                        return 'doc'
                    elif 'zip' in mime_type or 'x-zip' in mime_type:
                        return 'zip'
            except:
                pass  # Команда file может быть недоступна

            # МЕТОД 3: Fallback - по расширению (наименее надежный)
            ext = Path(file_path).suffix.lower()
            if ext == '.pdf':
                return 'pdf'
            elif ext in ['.xlsx', '.xls']:
                return 'xlsx'
            elif ext in ['.docx', '.doc']:
                return 'docx'
            elif ext == '.rtf':
                return 'rtf'
            elif ext == '.csv':
                return 'csv'
            elif ext in ['.txt', '.text', '.log']:
                return 'txt'

            return 'unknown'

        except Exception as e:
            print(f"⚠️  Ошибка определения типа файла {file_path}: {e}")
            # Последний fallback: используем расширение
            ext = Path(file_path).suffix.lower()
            if ext == '.pdf':
                return 'pdf'
            elif ext in ['.xlsx', '.xls']:
                return 'xlsx'
            elif ext in ['.docx', '.doc']:
                return 'docx'
            elif ext == '.rtf':
                return 'rtf'
            elif ext == '.csv':
                return 'csv'
            elif ext in ['.txt', '.text', '.log']:
                return 'txt'
            return 'unknown'

    @staticmethod
    def extract_from_pdf_with_ocr(file_path: str, max_pages: int = 20) -> str:
        """
        Извлекает текст из PDF используя OCR (для сканированных документов).

        Args:
            file_path: Путь к PDF файлу
            max_pages: Максимальное количество страниц для OCR (чтобы не перегружать систему)

        Returns:
            Извлеченный текст

        Raises:
            Exception: При ошибках OCR
        """
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from PIL import Image

            print(f"   🔍 Используем OCR для распознавания сканированного PDF...")

            # Конвертируем PDF в изображения
            try:
                images = convert_from_path(file_path, first_page=1, last_page=max_pages)
            except Exception as e:
                print(f"   ⚠️  Ошибка конвертации PDF в изображения: {e}")
                raise

            if not images:
                raise ValueError("Не удалось конвертировать PDF в изображения")

            print(f"   📄 Обработка {len(images)} страниц через OCR...")

            # Применяем OCR к каждому изображению
            text_content = []
            for i, image in enumerate(images, 1):
                try:
                    # Распознаем текст на русском и английском
                    text = pytesseract.image_to_string(image, lang='rus+eng')
                    if text.strip():
                        text_content.append(text.strip())
                    print(f"   ✓ Страница {i}/{len(images)} обработана")
                except Exception as page_error:
                    print(f"   ⚠️  Ошибка OCR на странице {i}: {page_error}")
                    continue

            extracted_text = '\n\n'.join(text_content)

            if not extracted_text.strip():
                raise ValueError("OCR не смог извлечь текст из PDF")

            print(f"   ✅ OCR завершен, извлечено {len(extracted_text)} символов")
            return extracted_text

        except ImportError as ie:
            missing_lib = str(ie).split("'")[1] if "'" in str(ie) else "библиотека"
            raise Exception(f"Для OCR требуется установить {missing_lib}: pip install pdf2image pytesseract pillow")
        except Exception as e:
            raise Exception(f"Ошибка OCR: {str(e)}")

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Извлекает текст из PDF файла.
        Сначала пытается извлечь обычным способом (PyPDF2),
        если не удается - использует OCR для сканированных документов.

        Args:
            file_path: Путь к PDF файлу

        Returns:
            Извлеченный текст

        Raises:
            FileNotFoundError: Если файл не найден
            Exception: При ошибках чтения PDF
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)

                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)

            extracted_text = '\n\n'.join(text_content)

            # Детальное логирование
            print(f"   📊 PDF Statistics:")
            print(f"      • Total pages: {total_pages}")
            print(f"      • Pages with text: {len(text_content)}")
            print(f"      • Total characters: {len(extracted_text):,}")
            print(f"      • Total words: {len(extracted_text.split()):,}")

            if not extracted_text.strip():
                # PDF не содержит текста - вероятно это скан
                # Пробуем OCR
                print(f"   ⚠️  PDF не содержит текстового слоя, пробуем OCR...")
                try:
                    return TextExtractor.extract_from_pdf_with_ocr(file_path)
                except Exception as ocr_error:
                    print(f"   ❌ OCR также не удался: {ocr_error}")
                    raise ValueError("PDF файл не содержит извлекаемого текста и OCR не помог")

            return extracted_text

        except Exception as e:
            # Если это не ошибка "нет текста", пробуем альтернативные методы
            if "не содержит извлекаемого текста" in str(e) or "EOF marker not found" in str(e):
                # СНАЧАЛА пробуем pdftotext (более надежный для битых PDF)
                try:
                    print(f"   ⚠️  PyPDF2 не смог извлечь текст ({str(e)}), пробуем pdftotext...")
                    result = subprocess.run(
                        ['pdftotext', '-layout', file_path, '-'],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        extracted_text = result.stdout.strip()
                        if len(extracted_text) > 100:  # Минимум 100 символов
                            print(f"   ✅ pdftotext успешно извлек {len(extracted_text)} символов")
                            return extracted_text
                except Exception as pdftotext_error:
                    print(f"   ⚠️  pdftotext также не помог: {pdftotext_error}")

                # Только если pdftotext не помог - пробуем OCR
                try:
                    print(f"   🔍 Пробуем OCR как последнюю попытку...")
                    return TextExtractor.extract_from_pdf_with_ocr(file_path)
                except Exception as ocr_error:
                    print(f"   ❌ OCR также не удался: {ocr_error}")

            raise Exception(f"Ошибка при извлечении текста из PDF: {str(e)}")

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Извлекает текст из DOCX файла.
        Поддерживает как новые (.docx), так и старые (.doc) форматы Word.

        Args:
            file_path: Путь к DOCX/DOC файлу

        Returns:
            Извлеченный текст

        Raises:
            FileNotFoundError: Если файл не найден
            Exception: При ошибках чтения DOCX
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        try:
            doc = Document(file_path)
            text_content = []

            # Извлекаем текст из параграфов
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    paragraph_count += 1

            # Извлекаем текст из таблиц
            table_count = 0
            table_rows_count = 0
            cells_extracted = 0

            for table in doc.tables:
                table_count += 1
                # Добавляем маркер начала таблицы для улучшения структуры
                text_content.append(f"\n=== ТАБЛИЦА {table_count} ===")

                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        # УЛУЧШЕНИЕ 1: Извлекаем текст из всех параграфов внутри ячейки
                        # python-docx иногда не видит cell.text если там есть форматирование
                        cell_content = []
                        for paragraph in cell.paragraphs:
                            para_text = paragraph.text.strip()
                            if para_text:
                                cell_content.append(para_text)

                        # УЛУЧШЕНИЕ 2: Извлекаем текст из вложенных таблиц внутри ячейки
                        # Некоторые DOCX используют таблицы внутри таблиц для форматирования
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                nested_row_text = []
                                for nested_cell in nested_row.cells:
                                    nested_text = nested_cell.text.strip()
                                    if nested_text:
                                        nested_row_text.append(nested_text)
                                if nested_row_text:
                                    cell_content.append(' '.join(nested_row_text))

                        if cell_content:
                            # Объединяем параграфы внутри ячейки через пробел
                            combined = ' '.join(cell_content)
                            row_text.append(combined)
                            cells_extracted += 1

                    if row_text:
                        text_content.append(' | '.join(row_text))
                        table_rows_count += 1

            extracted_text = '\n\n'.join(text_content)

            # Детальное логирование
            print(f"   📊 DOCX Statistics:")
            print(f"      • Paragraphs extracted: {paragraph_count}")
            print(f"      • Tables found: {table_count}")
            print(f"      • Table rows extracted: {table_rows_count}")
            print(f"      • Table cells extracted: {cells_extracted}")
            print(f"      • Total characters: {len(extracted_text):,}")
            print(f"      • Total words: {len(extracted_text.split()):,}")

            # ДИАГНОСТИКА: Показываем первые 1000 символов из таблиц
            if table_count > 0:
                # Ищем текст после первого маркера таблицы
                table_start = extracted_text.find("=== ТАБЛИЦА 1 ===")
                if table_start != -1:
                    table_preview = extracted_text[table_start:table_start + 1000]
                    print(f"      • Table preview (first 1000 chars):")
                    # Показываем построчно для читаемости
                    for line in table_preview.split('\n')[:10]:
                        if line.strip():
                            print(f"        {line[:100]}")

            if not extracted_text.strip():
                raise ValueError("DOCX файл не содержит текста")

            return extracted_text

        except Exception as e:
            # Fallback: пытаемся извлечь текст через antiword для старых .doc файлов
            try:
                print(f"   ⚠️  Ошибка python-docx, пытаемся использовать antiword...")
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    print(f"   ✅ Текст извлечен через antiword")
                    return result.stdout.strip()
            except Exception as antiword_error:
                print(f"   ⚠️  antiword тоже не помог: {antiword_error}")

            # Fallback 2: пытаемся textract (если установлен)
            try:
                import textract
                print(f"   ⚠️  Пытаемся использовать textract...")
                text = textract.process(file_path).decode('utf-8')
                if text.strip():
                    print(f"   ✅ Текст извлечен через textract")
                    return text.strip()
            except ImportError:
                pass
            except Exception as textract_error:
                print(f"   ⚠️  textract не помог: {textract_error}")

            # Если все попытки провалились, возвращаем частичную информацию
            print(f"   ❌ Не удалось извлечь текст из документа")
            raise Exception(f"Ошибка при извлечении текста из DOCX: {str(e)}")

    @staticmethod
    def extract_from_xlsx(file_path: str) -> str:
        """
        Извлекает текст из XLSX/XLS файла (Excel).

        Args:
            file_path: Путь к XLSX/XLS файлу

        Returns:
            Извлеченный текст

        Raises:
            FileNotFoundError: Если файл не найден
            Exception: При ошибках чтения XLSX
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        try:
            # Загружаем книгу Excel
            workbook = load_workbook(file_path, data_only=True)
            text_content = []
            total_sheets = 0
            total_rows = 0

            # Обрабатываем каждый лист
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                total_sheets += 1

                # Добавляем заголовок с именем листа
                text_content.append(f"=== Лист: {sheet_name} ===")

                # Извлекаем данные из всех ячеек
                sheet_rows = []
                for row in sheet.iter_rows(values_only=True):
                    # Фильтруем пустые ячейки и конвертируем в строку
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    # Убираем полностью пустые строки
                    if any(val.strip() for val in row_values):
                        sheet_rows.append(' | '.join(row_values))

                if sheet_rows:
                    text_content.append('\n'.join(sheet_rows))
                    total_rows += len(sheet_rows)

            extracted_text = '\n\n'.join(text_content)

            # Детальное логирование
            print(f"   📊 XLSX Statistics:")
            print(f"      • Sheets processed: {total_sheets}")
            print(f"      • Rows extracted: {total_rows}")
            print(f"      • Total characters: {len(extracted_text):,}")
            print(f"      • Total words: {len(extracted_text.split()):,}")

            if not extracted_text.strip():
                raise ValueError("XLSX файл не содержит текста")

            return extracted_text

        except Exception as e:
            print(f"   ❌ Ошибка извлечения текста из XLSX: {e}")
            raise Exception(f"Ошибка при извлечении текста из XLSX: {str(e)}")

    @staticmethod
    def extract_from_text_file(file_path: str) -> str:
        """
        Извлекает текст из простого текстового файла (TXT, LOG и т.д.).
        Автоматически определяет кодировку.

        Args:
            file_path: Путь к текстовому файлу

        Returns:
            Извлеченный текст

        Raises:
            FileNotFoundError: Если файл не найден
            Exception: При ошибках чтения файла
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        try:
            # Определяем кодировку файла
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected = chardet.detect(raw_data)
                encoding = detected['encoding'] or 'utf-8'

            # Читаем файл с определенной кодировкой
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                text = file.read()

            if not text.strip():
                raise ValueError("Текстовый файл пустой")

            return text

        except Exception as e:
            print(f"   ❌ Ошибка извлечения текста из TXT: {e}")
            raise Exception(f"Ошибка при извлечении текста из текстового файла: {str(e)}")

    @staticmethod
    def extract_from_csv(file_path: str) -> str:
        """
        Извлекает текст из CSV файла.

        Args:
            file_path: Путь к CSV файлу

        Returns:
            Извлеченный текст в табличном формате

        Raises:
            FileNotFoundError: Если файл не найден
            Exception: При ошибках чтения CSV
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        try:
            # Определяем кодировку файла
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected = chardet.detect(raw_data)
                encoding = detected['encoding'] or 'utf-8'

            # Читаем CSV
            text_content = []
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                csv_reader = csv.reader(file)
                for row_num, row in enumerate(csv_reader, 1):
                    # Фильтруем пустые строки
                    if any(cell.strip() for cell in row):
                        text_content.append(' | '.join(row))

            extracted_text = '\n'.join(text_content)

            if not extracted_text.strip():
                raise ValueError("CSV файл пустой")

            return extracted_text

        except Exception as e:
            print(f"   ❌ Ошибка извлечения текста из CSV: {e}")
            raise Exception(f"Ошибка при извлечении текста из CSV: {str(e)}")

    @staticmethod
    def extract_from_rtf(file_path: str) -> str:
        """
        Извлекает текст из RTF файла.
        Использует striprtf если доступен, иначе простое чтение.

        Args:
            file_path: Путь к RTF файлу

        Returns:
            Извлеченный текст

        Raises:
            FileNotFoundError: Если файл не найден
            Exception: При ошибках чтения RTF
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        try:
            # Пробуем использовать striprtf для качественного извлечения
            try:
                from striprtf.striprtf import rtf_to_text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    rtf_content = file.read()
                    text = rtf_to_text(rtf_content)
                    if text and text.strip():
                        return text
            except ImportError:
                print(f"   ⚠️  striprtf не установлен, используем базовое извлечение")
            except Exception as e:
                print(f"   ⚠️  Ошибка striprtf ({e}), используем базовое извлечение")

            # Fallback: используем unrtf через subprocess
            try:
                result = subprocess.run(
                    ['unrtf', '--text', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except Exception as e:
                print(f"   ⚠️  unrtf не доступен: {e}")

            # Последний fallback: просто читаем как текст (будет с RTF разметкой)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                if text.strip():
                    print(f"   ⚠️  Используем сырое чтение RTF (может содержать разметку)")
                    return text

            raise ValueError("Не удалось извлечь текст из RTF файла")

        except Exception as e:
            print(f"   ❌ Ошибка извлечения текста из RTF: {e}")
            raise Exception(f"Ошибка при извлечении текста из RTF: {str(e)}")

    @staticmethod
    def extract_from_unknown(file_path: str) -> str:
        """
        Универсальный fallback для неизвестных типов файлов.
        Пытается извлечь текст любыми доступными способами.

        Args:
            file_path: Путь к файлу

        Returns:
            Извлеченный текст или сообщение о невозможности извлечения

        Raises:
            Exception: При критических ошибках
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        print(f"   🔧 Попытка универсального извлечения текста...")

        # 1. Пробуем прочитать как текст с автоопределением кодировки
        try:
            with open(file_path, 'rb') as file:
                raw_data = file.read(10000)  # Читаем первые 10KB для определения кодировки
                detected = chardet.detect(raw_data)
                encoding = detected['encoding'] or 'utf-8'
                confidence = detected['confidence']

            if confidence > 0.7:  # Достаточная уверенность в кодировке
                with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                    text = file.read()
                    # Проверяем, что текст содержит хотя бы некоторые читаемые символы
                    readable_chars = sum(1 for c in text if c.isprintable() or c.isspace())
                    if readable_chars > len(text) * 0.5:  # >50% читаемых символов
                        print(f"   ✅ Текст извлечен как {encoding} (уверенность: {confidence:.0%})")
                        return text
        except Exception as e:
            print(f"   ⚠️  Попытка чтения как текст не удалась: {e}")

        # 2. Пробуем strings (извлечение текстовых строк из бинарного файла)
        # ТОЛЬКО для неповрежденных документов
        try:
            result = subprocess.run(
                ['strings', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                extracted = result.stdout.strip()

                # Валидация: проверяем что это не мусор
                # Считаем процент русских/английских букв от всех символов
                letters = sum(1 for c in extracted if c.isalpha())
                total_printable = sum(1 for c in extracted if c.isprintable())

                if total_printable > 100 and letters / total_printable > 0.4:
                    # Хотя бы 40% букв - вероятно это текст
                    print(f"   ✅ Извлечены текстовые строки из бинарного файла")
                    return extracted
                else:
                    print(f"   ⚠️  Извлеченные строки не похожи на текст (букв: {letters}/{total_printable})")
        except Exception as e:
            print(f"   ⚠️  Команда strings не доступна: {e}")

        # 3. Последняя попытка: читаем как UTF-8 с игнорированием ошибок
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                if text.strip():
                    print(f"   ⚠️  Использовано принудительное чтение как UTF-8")
                    return text
        except Exception as e:
            print(f"   ❌ Все попытки извлечения не удались: {e}")

        # Если ничего не помогло, возвращаем информативное сообщение
        file_size = os.path.getsize(file_path)
        ext = Path(file_path).suffix
        return f"[Не удалось автоматически извлечь текст из файла {ext} размером {file_size} байт. Возможно, это бинарный файл или неподдерживаемый формат.]"

    @staticmethod
    def extract_from_zip(file_path: str) -> str:
        """
        Извлекает текст из ZIP-архива, распаковывая его содержимое.
        Автоматически определяет DOCX файлы (которые тоже являются ZIP).

        Args:
            file_path: Путь к ZIP файлу

        Returns:
            Извлеченный текст из всех поддерживаемых файлов в архиве
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        try:
            # Сначала проверяем, не является ли это DOCX файлом
            # DOCX это ZIP архив с определенной структурой (содержит word/document.xml)
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()

                # Проверяем признаки DOCX/Office Open XML
                if 'word/document.xml' in file_list or '[Content_Types].xml' in file_list:
                    print(f"   📦 Обнаружен Office документ (DOCX) внутри ZIP")
                    # Это DOCX файл, используем python-docx для извлечения
                    try:
                        doc = Document(file_path)
                        text_content = []

                        # Извлекаем текст из параграфов
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                text_content.append(paragraph.text)

                        # Извлекаем текст из таблиц
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    if cell.text.strip():
                                        row_text.append(cell.text.strip())
                                if row_text:
                                    text_content.append(' | '.join(row_text))

                        extracted_text = '\n\n'.join(text_content)

                        if not extracted_text.strip():
                            raise ValueError("DOCX файл не содержит текста")

                        return extracted_text
                    except Exception as e:
                        raise Exception(f"Ошибка извлечения текста из DOCX в ZIP: {str(e)}")

            # Если это обычный ZIP с файлами (не DOCX)
            with tempfile.TemporaryDirectory() as temp_dir:
                # Распаковываем архив
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)

                # Получаем список всех файлов (исключая служебные)
                extracted_files = []
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        # Пропускаем XML и служебные файлы DOCX
                        if (not file.startswith('.') and
                            not file.startswith('__') and
                            not file.endswith('.xml') and
                            not file.endswith('.rels') and
                            not file.endswith('.bin')):
                            extracted_files.append(os.path.join(root, file))

                if not extracted_files:
                    # Попробуем все файлы, включая XML
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            if not file.startswith('.') and not file.startswith('__'):
                                extracted_files.append(os.path.join(root, file))

                if not extracted_files:
                    raise ValueError("ZIP архив пустой или содержит только служебные файлы")

                print(f"   📦 Найдено файлов в архиве: {len(extracted_files)}")

                # Извлекаем текст из каждого файла
                all_texts = []
                for extracted_file in extracted_files:
                    try:
                        # Рекурсивно вызываем extract_text для каждого файла
                        result = TextExtractor.extract_text(extracted_file)
                        if result['text'] and not result['text'].startswith('[Не удалось'):
                            all_texts.append(f"=== {result['file_name']} ===\n{result['text']}")
                    except Exception as e:
                        # Просто пропускаем файлы, которые не удалось обработать
                        continue

                if not all_texts:
                    raise ValueError("Не удалось извлечь текст ни из одного файла в архиве")

                return '\n\n'.join(all_texts)

        except zipfile.BadZipFile:
            raise Exception("Файл поврежден или не является корректным ZIP архивом")
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из ZIP: {str(e)}")

    @staticmethod
    def extract_text(file_path: str) -> dict:
        """
        Универсальный метод для извлечения текста из поддерживаемых форматов.

        Args:
            file_path: Путь к файлу

        Returns:
            Словарь с извлеченным текстом и метаданными:
            {
                'text': str,
                'file_name': str,
                'file_type': str,
                'char_count': int,
                'word_count': int
            }

        Raises:
            ValueError: Если формат не поддерживается
            Exception: При ошибках извлечения
        """
        file_path = str(Path(file_path).resolve())

        # ВАЖНО: Определяем реальный тип файла по содержимому, а не по расширению!
        # Zakupki.gov.ru часто называет Word файлы с расширением .pdf
        actual_type = TextExtractor.detect_file_type(file_path)

        print(f"   📄 Файл: {Path(file_path).name}")
        print(f"   🔍 Определен тип: {actual_type}")

        # Определяем метод извлечения по РЕАЛЬНОМУ типу файла
        try:
            if actual_type == 'pdf':
                text = TextExtractor.extract_from_pdf(file_path)
                file_type = 'PDF'
            elif actual_type in ['docx', 'doc']:
                text = TextExtractor.extract_from_docx(file_path)
                file_type = 'DOCX/DOC'
            elif actual_type == 'xlsx':
                text = TextExtractor.extract_from_xlsx(file_path)
                file_type = 'XLSX/XLS'
            elif actual_type == 'txt':
                text = TextExtractor.extract_from_text_file(file_path)
                file_type = 'TXT'
            elif actual_type == 'csv':
                text = TextExtractor.extract_from_csv(file_path)
                file_type = 'CSV'
            elif actual_type == 'rtf':
                text = TextExtractor.extract_from_rtf(file_path)
                file_type = 'RTF'
            elif actual_type == 'zip':
                text = TextExtractor.extract_from_zip(file_path)
                file_type = 'ZIP'
            else:
                # Используем универсальный fallback для неизвестных типов
                file_extension = Path(file_path).suffix.lower()
                print(f"   ⚠️  Неизвестный формат: {actual_type}, пробуем универсальное извлечение")
                text = TextExtractor.extract_from_unknown(file_path)
                file_type = f'UNKNOWN ({file_extension})'
        except Exception as extract_error:
            # Если не удалось извлечь текст стандартными методами, пробуем универсальный fallback
            print(f"   ❌ Ошибка извлечения текста: {extract_error}")
            try:
                print(f"   🔄 Пробуем универсальный метод извлечения...")
                text = TextExtractor.extract_from_unknown(file_path)
                file_type = f'{actual_type.upper()} (fallback)'
            except Exception as fallback_error:
                print(f"   ❌ Универсальное извлечение также не удалось: {fallback_error}")
                text = f"[Не удалось извлечь текст из файла: {str(extract_error)[:200]}]"
                file_type = actual_type.upper()

        # Подсчитываем статистику
        char_count = len(text)
        word_count = len(text.split())

        # Итоговое логирование результата
        print(f"   ✅ Extraction completed:")
        print(f"      • File: {os.path.basename(file_path)}")
        print(f"      • Type: {file_type}")
        print(f"      • Characters: {char_count:,}")
        print(f"      • Words: {word_count:,}")
        if char_count > 0:
            # Показываем первые 200 символов для проверки
            preview = text[:200].replace('\n', ' ')
            print(f"      • Preview: {preview}...")

        return {
            'text': text,
            'file_name': os.path.basename(file_path),
            'file_type': file_type,
            'char_count': char_count,
            'word_count': word_count
        }

    @staticmethod
    def extract_from_multiple_files(file_paths: list) -> dict:
        """
        Извлекает текст из нескольких файлов и объединяет их.

        Args:
            file_paths: Список путей к файлам

        Returns:
            Словарь с объединенным текстом и метаданными:
            {
                'combined_text': str,
                'files': list[dict],
                'total_char_count': int,
                'total_word_count': int
            }
        """
        results = []
        combined_text = []

        for file_path in file_paths:
            try:
                result = TextExtractor.extract_text(file_path)
                results.append(result)
                combined_text.append(f"=== {result['file_name']} ===\n{result['text']}")
            except Exception as e:
                results.append({
                    'file_name': os.path.basename(file_path),
                    'error': str(e)
                })

        full_text = '\n\n'.join(combined_text)

        return {
            'combined_text': full_text,
            'files': results,
            'total_char_count': len(full_text),
            'total_word_count': len(full_text.split())
        }


def main():
    """Пример использования TextExtractor."""
    import sys

    if len(sys.argv) < 2:
        print("Использование: python text_extractor.py <путь_к_файлу>")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        result = TextExtractor.extract_text(file_path)
        print(f"\n{'='*60}")
        print(f"Файл: {result['file_name']}")
        print(f"Тип: {result['file_type']}")
        print(f"Символов: {result['char_count']:,}")
        print(f"Слов: {result['word_count']:,}")
        print(f"{'='*60}\n")
        print(result['text'][:500] + "..." if len(result['text']) > 500 else result['text'])
    except Exception as e:
        print(f"Ошибка: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
