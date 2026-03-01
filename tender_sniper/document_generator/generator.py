"""
DocumentGenerator — основной класс генерации тендерных документов.

Заполняет DOCX-шаблоны данными компании + тендера.
Документы генерируются на лету (Railway — ephemeral FS).
"""

import io
import os
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent / 'templates'

# Типы документов
DOC_TYPES = {
    'application': 'Заявка на участие',
    'declaration': 'Декларация соответствия',
    'agreement': 'Согласие с условиями',
    'proposal': 'Техническое предложение',
}


class DocumentGenerator:
    """Генератор пакета тендерных документов."""

    def __init__(self):
        self.templates_dir = TEMPLATES_DIR

    async def generate_package(
        self,
        tender_data: Dict[str, Any],
        company_profile: Dict[str, Any],
        user_id: int,
        ai_proposal_text: Optional[str] = None,
    ) -> List[Tuple[str, str, io.BytesIO]]:
        """
        Генерация полного пакета документов.

        Args:
            tender_data: Данные тендера (number, name, price, url, customer_name, etc.)
            company_profile: Профиль компании из БД
            user_id: ID пользователя
            ai_proposal_text: Предгенерированный AI-текст техпредложения

        Returns:
            Список кортежей: (doc_type, filename, BytesIO с DOCX)
        """
        context = self._build_context(tender_data, company_profile)
        documents = []

        for doc_type in ['application', 'declaration', 'agreement', 'proposal']:
            try:
                if doc_type == 'proposal' and ai_proposal_text:
                    context['proposal_text'] = ai_proposal_text

                doc_bytes = self._generate_document(doc_type, context)
                tender_num = tender_data.get('number', 'unknown')[:30]
                filename = f"{DOC_TYPES[doc_type]}_{tender_num}.docx"
                # Очистка имени файла
                filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                documents.append((doc_type, filename, doc_bytes))
                logger.info(f"Generated {doc_type} for tender {tender_num}")
            except Exception as e:
                logger.error(f"Error generating {doc_type}: {e}", exc_info=True)

        return documents

    def _build_context(self, tender_data: Dict[str, Any], profile: Dict[str, Any]) -> Dict[str, str]:
        """Построение контекста подстановки для шаблонов."""
        now = datetime.utcnow()

        return {
            # Компания
            'company_name': profile.get('company_name', '_______________'),
            'company_name_short': profile.get('company_name_short') or profile.get('company_name', '_______________'),
            'legal_form': profile.get('legal_form', ''),
            'inn': profile.get('inn', '_______________'),
            'kpp': profile.get('kpp', ''),
            'ogrn': profile.get('ogrn', ''),
            'legal_address': profile.get('legal_address', '_______________'),
            'actual_address': profile.get('actual_address') or profile.get('legal_address', '_______________'),
            'postal_address': profile.get('postal_address') or profile.get('legal_address', '_______________'),
            'director_name': profile.get('director_name', '_______________'),
            'director_position': profile.get('director_position', 'Генеральный директор'),
            'director_basis': profile.get('director_basis', 'Устав'),
            'phone': profile.get('phone', '_______________'),
            'email': profile.get('email', '_______________'),
            'website': profile.get('website', ''),
            'bank_name': profile.get('bank_name', ''),
            'bank_bik': profile.get('bank_bik', ''),
            'bank_account': profile.get('bank_account', ''),
            'bank_corr_account': profile.get('bank_corr_account', ''),
            'smp_status': 'Да' if profile.get('smp_status') else 'Нет',
            'licenses_text': profile.get('licenses_text', 'Не требуются'),
            'experience_description': profile.get('experience_description', ''),

            # Тендер
            'tender_number': tender_data.get('number', ''),
            'tender_name': tender_data.get('name', ''),
            'tender_price': self._format_price(tender_data.get('price')),
            'tender_url': tender_data.get('url', ''),
            'tender_customer': tender_data.get('customer_name', ''),
            'tender_region': tender_data.get('region', ''),
            'tender_deadline': tender_data.get('submission_deadline', ''),

            # Дата
            'current_date': now.strftime('%d.%m.%Y'),
            'current_year': str(now.year),

            # Для техпредложения
            'proposal_text': '',
        }

    def _format_price(self, price) -> str:
        """Форматирование цены."""
        if price is None:
            return 'не указана'
        try:
            p = float(price)
            if p >= 1_000_000:
                return f"{p/1_000_000:.2f} млн руб."
            elif p >= 1_000:
                return f"{p/1_000:.1f} тыс. руб."
            else:
                return f"{p:.2f} руб."
        except (ValueError, TypeError):
            return str(price)

    def _generate_document(self, doc_type: str, context: Dict[str, str]) -> io.BytesIO:
        """
        Генерация одного документа из шаблона.

        Если шаблон DOCX существует — использует его.
        Иначе — генерирует программно через python-docx.
        """
        template_path = self.templates_dir / f"{doc_type}.docx"

        if template_path.exists():
            return self._fill_template(template_path, context)
        else:
            return self._generate_programmatic(doc_type, context)

    def _fill_template(self, template_path: Path, context: Dict[str, str]) -> io.BytesIO:
        """Заполнение DOCX-шаблона подстановкой {{placeholder}}."""
        doc = Document(str(template_path))

        # Замена в параграфах
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, context)

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, context)

        # Замена в headers/footers
        for section in doc.sections:
            for header in [section.header, section.footer]:
                if header:
                    for paragraph in header.paragraphs:
                        self._replace_in_paragraph(paragraph, context)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _replace_in_paragraph(self, paragraph, context: Dict[str, str]):
        """Замена плейсхолдеров в параграфе с сохранением форматирования."""
        full_text = paragraph.text
        if '{{' not in full_text:
            return

        for key, value in context.items():
            placeholder = '{{' + key + '}}'
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value or ''))

        # Сохраняем форматирование первого run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = full_text
                else:
                    run.text = ''

    def _generate_programmatic(self, doc_type: str, ctx: Dict[str, str]) -> io.BytesIO:
        """Программная генерация документа (если нет шаблона)."""
        generators = {
            'application': self._gen_application,
            'declaration': self._gen_declaration,
            'agreement': self._gen_agreement,
            'proposal': self._gen_proposal,
        }

        generator = generators.get(doc_type)
        if not generator:
            raise ValueError(f"Unknown doc_type: {doc_type}")

        doc = Document()

        # Настройка шрифта по умолчанию (Times New Roman — стандарт для госдокументов)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Настройка полей страницы
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)

        generator(doc, ctx)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _add_header(self, doc: Document, text: str):
        """Добавление заголовка."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False,
                       align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """Добавление параграфа."""
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def _gen_application(self, doc: Document, ctx: Dict[str, str]):
        """Генерация заявки на участие."""
        # Шапка
        self._add_paragraph(doc, f"Заказчику: {ctx['tender_customer']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        doc.add_paragraph()

        self._add_header(doc, "ЗАЯВКА НА УЧАСТИЕ В ЗАКУПКЕ")
        doc.add_paragraph()

        self._add_paragraph(doc, f"Закупка № {ctx['tender_number']}")
        self._add_paragraph(doc, f"«{ctx['tender_name']}»")
        doc.add_paragraph()

        # Данные участника
        self._add_paragraph(doc, "1. Сведения об участнике закупки:", bold=True)
        self._add_paragraph(doc, f"Наименование: {ctx['company_name']}")
        self._add_paragraph(doc, f"ИНН: {ctx['inn']}")
        if ctx.get('kpp'):
            self._add_paragraph(doc, f"КПП: {ctx['kpp']}")
        if ctx.get('ogrn'):
            self._add_paragraph(doc, f"ОГРН: {ctx['ogrn']}")
        self._add_paragraph(doc, f"Юридический адрес: {ctx['legal_address']}")
        self._add_paragraph(doc, f"Фактический адрес: {ctx['actual_address']}")
        self._add_paragraph(doc, f"Телефон: {ctx['phone']}")
        self._add_paragraph(doc, f"Электронная почта: {ctx['email']}")

        doc.add_paragraph()
        self._add_paragraph(doc, "2. Банковские реквизиты:", bold=True)
        if ctx.get('bank_name'):
            self._add_paragraph(doc, f"Банк: {ctx['bank_name']}")
        if ctx.get('bank_bik'):
            self._add_paragraph(doc, f"БИК: {ctx['bank_bik']}")
        if ctx.get('bank_account'):
            self._add_paragraph(doc, f"Расчётный счёт: {ctx['bank_account']}")
        if ctx.get('bank_corr_account'):
            self._add_paragraph(doc, f"Корреспондентский счёт: {ctx['bank_corr_account']}")

        doc.add_paragraph()
        self._add_paragraph(doc, "3. Предложение участника:", bold=True)
        self._add_paragraph(
            doc,
            f"Изучив закупочную документацию по закупке № {ctx['tender_number']} "
            f"«{ctx['tender_name']}», {ctx['company_name']} выражает готовность "
            f"выполнить условия закупки в полном объёме."
        )

        doc.add_paragraph()
        self._add_paragraph(doc, f"Является субъектом МСП: {ctx['smp_status']}")

        # Подпись
        doc.add_paragraph()
        doc.add_paragraph()
        self._add_paragraph(doc, f"{ctx['director_position']}")
        self._add_paragraph(doc, f"_________________ / {ctx['director_name']}")
        self._add_paragraph(doc, f"М.П.")
        self._add_paragraph(doc, f"«___» ____________ {ctx['current_year']} г.")

    def _gen_declaration(self, doc: Document, ctx: Dict[str, str]):
        """Генерация декларации соответствия требованиям."""
        self._add_header(doc, "ДЕКЛАРАЦИЯ")
        self._add_header(doc, "о соответствии участника закупки требованиям,")
        self._add_header(doc, "установленным в документации о закупке")
        doc.add_paragraph()

        self._add_paragraph(
            doc,
            f"{ctx['company_name']} (ИНН {ctx['inn']}), "
            f"в лице {ctx['director_position']} {ctx['director_name']}, "
            f"действующего на основании {ctx['director_basis']}, настоящим заявляет и декларирует следующее:"
        )

        doc.add_paragraph()
        points = [
            "Участник закупки не является иностранным юридическим лицом, не является российским юридическим лицом, "
            "в уставном (складочном) капитале которого доля участия иностранных лиц превышает 25 процентов.",
            "Участник закупки не является офшорной компанией.",
            "На участника закупки не распространяются ограничения для участия в закупках, установленные законодательством РФ.",
            "Отсутствует решение арбитражного суда о признании участника закупки банкротом.",
            "Деятельность участника закупки не приостановлена в порядке, предусмотренном КоАП РФ.",
            "У участника закупки отсутствует задолженность по налогам, сборам и иным обязательным платежам за прошедший год, "
            "превышающая 25% активов участника закупки.",
            "У участника закупки отсутствует судимость за преступления в сфере экономики у руководителя, "
            "членов коллегиального исполнительного органа, главного бухгалтера.",
            "Между участником закупки и заказчиком отсутствует конфликт интересов.",
        ]

        for i, point in enumerate(points, 1):
            self._add_paragraph(doc, f"{i}. {point}")

        doc.add_paragraph()
        self._add_paragraph(doc, f"Дата: {ctx['current_date']}")
        doc.add_paragraph()
        self._add_paragraph(doc, f"{ctx['director_position']}")
        self._add_paragraph(doc, f"_________________ / {ctx['director_name']}")
        self._add_paragraph(doc, "М.П.")

    def _gen_agreement(self, doc: Document, ctx: Dict[str, str]):
        """Генерация согласия с условиями закупки."""
        self._add_paragraph(doc, f"Заказчику: {ctx['tender_customer']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        doc.add_paragraph()

        self._add_header(doc, "СОГЛАСИЕ")
        self._add_header(doc, "с условиями исполнения контракта")
        doc.add_paragraph()

        self._add_paragraph(
            doc,
            f"{ctx['company_name']} (ИНН {ctx['inn']}), "
            f"в лице {ctx['director_position']} {ctx['director_name']}, "
            f"действующего на основании {ctx['director_basis']}, "
            f"настоящим подтверждает согласие с условиями исполнения контракта, "
            f"указанными в документации о закупке № {ctx['tender_number']} "
            f"«{ctx['tender_name']}»."
        )

        doc.add_paragraph()
        self._add_paragraph(doc, "В том числе подтверждаем:")
        doc.add_paragraph()

        items = [
            "Согласие на выполнение работ/оказание услуг/поставку товаров на условиях, "
            "указанных в документации о закупке.",
            "Согласие с порядком оплаты, указанным в проекте контракта.",
            "Согласие на выполнение обязательств в сроки, установленные документацией о закупке.",
            "Готовность предоставить обеспечение исполнения контракта в случае, "
            "если такое требование предусмотрено документацией о закупке.",
        ]

        for i, item in enumerate(items, 1):
            self._add_paragraph(doc, f"{i}. {item}")

        doc.add_paragraph()
        self._add_paragraph(doc, f"Дата: {ctx['current_date']}")
        doc.add_paragraph()
        self._add_paragraph(doc, f"{ctx['director_position']}")
        self._add_paragraph(doc, f"_________________ / {ctx['director_name']}")
        self._add_paragraph(doc, "М.П.")

    def _gen_proposal(self, doc: Document, ctx: Dict[str, str]):
        """Генерация технического предложения."""
        self._add_paragraph(doc, f"Заказчику: {ctx['tender_customer']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        doc.add_paragraph()

        self._add_header(doc, "ТЕХНИЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
        doc.add_paragraph()

        self._add_paragraph(doc, f"по закупке № {ctx['tender_number']}")
        self._add_paragraph(doc, f"«{ctx['tender_name']}»")
        doc.add_paragraph()

        self._add_paragraph(doc, "1. Сведения об участнике:", bold=True)
        self._add_paragraph(doc, f"Наименование: {ctx['company_name']}")
        self._add_paragraph(doc, f"ИНН: {ctx['inn']}")
        self._add_paragraph(doc, f"Адрес: {ctx['legal_address']}")

        if ctx.get('experience_description'):
            doc.add_paragraph()
            self._add_paragraph(doc, "2. Опыт работы:", bold=True)
            self._add_paragraph(doc, ctx['experience_description'])

        if ctx.get('licenses_text') and ctx['licenses_text'] != 'Не требуются':
            doc.add_paragraph()
            self._add_paragraph(doc, "3. Лицензии и допуски:", bold=True)
            self._add_paragraph(doc, ctx['licenses_text'])

        # AI-сгенерированное техпредложение
        proposal_text = ctx.get('proposal_text', '')
        if proposal_text:
            doc.add_paragraph()
            self._add_paragraph(doc, "4. Техническое предложение:", bold=True)
            # Разбиваем по абзацам
            for para_text in proposal_text.split('\n'):
                para_text = para_text.strip()
                if para_text:
                    self._add_paragraph(doc, para_text)
        else:
            doc.add_paragraph()
            self._add_paragraph(doc, "4. Техническое предложение:", bold=True)
            self._add_paragraph(
                doc,
                f"{ctx['company_name']} готово выполнить работы/оказать услуги/осуществить поставку "
                f"в полном соответствии с требованиями закупочной документации."
            )

        # Подпись
        doc.add_paragraph()
        doc.add_paragraph()
        self._add_paragraph(doc, f"{ctx['director_position']}")
        self._add_paragraph(doc, f"_________________ / {ctx['director_name']}")
        self._add_paragraph(doc, "М.П.")
        self._add_paragraph(doc, f"«___» ____________ {ctx['current_year']} г.")
